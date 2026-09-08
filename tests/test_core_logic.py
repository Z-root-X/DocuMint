import os
import sys
from docx import Document
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))

try:
    from documint.core import is_valid_email, _replace_in_paragraph, replace_placeholders_in_doc, validate_placeholders
except ImportError:
    from src.documint.core import is_valid_email, _replace_in_paragraph, replace_placeholders_in_doc, validate_placeholders

def test_manual_string_replacement():
    """Verify that manual string replacement works as expected."""
    replacements = {"<Name>": "Zihad", "<ID>": "123"}
    body = "Hello <Name>, your ID is <ID>."
    
    final_body = body
    for k, v in replacements.items():
        final_body = final_body.replace(k, v)
        
    assert final_body == "Hello Zihad, your ID is 123."

def test_missing_placeholder():
    """Verify behavior when a placeholder is missing in replacements."""
    replacements = {"<Name>": "Zihad"}
    body = "Hello <Name>, your ID is <ID>."
    
    final_body = body
    for k, v in replacements.items():
        final_body = final_body.replace(k, v)
        
    # <ID> should remain untouched
    assert final_body == "Hello Zihad, your ID is <ID>."

def test_email_validation():
    """Test valid and invalid email addresses."""
    assert is_valid_email("user@example.com") is True
    assert is_valid_email("first.last@domain.co.uk") is True
    assert is_valid_email("invalid-email") is False
    assert is_valid_email("") is False
    assert is_valid_email("@domain.com") is False

def test_replace_placeholders_in_docx():
    """Test docx placeholder replacement across paragraph runs and curly/angle bracket syntaxes."""
    doc = Document()
    p1 = doc.add_paragraph("Welcome <Name>, to {{Course}}.")
    p2 = doc.add_paragraph("Your score is [Score].")
    
    replacements = {
        "Name": "Zihad Hasan",
        "Course": "Deep Learning",
        "Score": "98%"
    }
    
    replace_placeholders_in_doc(doc, replacements)
    
    assert doc.paragraphs[0].text == "Welcome Zihad Hasan, to Deep Learning."
    assert doc.paragraphs[1].text == "Your score is 98%."

def test_run_formatting_preservation():
    """Verify that replacing a placeholder inside a formatted run preserves run boldness."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("<Participant>")
    run.bold = True
    
    replacements = {"<Participant>": "Zihad Hasan"}
    replace_placeholders_in_doc(doc, replacements)
    
    assert p.text == "Zihad Hasan"
    assert p.runs[0].bold is True

def test_validate_placeholders(tmp_path):
    """Verify placeholder validation against data columns."""
    import pandas as pd
    
    # Create temp csv
    csv_file = tmp_path / "data.csv"
    df = pd.DataFrame({"Name": ["Zihad"], "Score": [99]})
    df.to_csv(csv_file, index=False)
    
    # Create temp docx with matching placeholders
    doc_file = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Hello <Name>, your score is {{Score}}.")
    doc.save(str(doc_file))
    
    is_valid, missing = validate_placeholders(str(csv_file), str(doc_file))
    assert is_valid is True
    assert missing == []
    
    # Create temp docx with missing placeholder
    doc_missing = tmp_path / "template_missing.docx"
    doc2 = Document()
    doc2.add_paragraph("Hello <Name>, your score is {{Score}} for course <MissingCourse>.")
    doc2.save(str(doc_missing))
    
    is_valid_bad, missing_bad = validate_placeholders(str(csv_file), str(doc_missing))
    assert is_valid_bad is False
    assert "MissingCourse" in missing_bad

