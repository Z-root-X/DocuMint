import pandas as pd
from docx import Document
import os

# Ensure examples directory exists
os.makedirs('examples', exist_ok=True)

# 1. Create Excel Template
df = pd.DataFrame([
    {'Name': 'John Doe', 'ID': 'EMP001', 'Department': 'Engineering', 'Email': 'john@example.com'},
    {'Name': 'Jane Smith', 'ID': 'EMP002', 'Department': 'HR', 'Email': 'jane@example.com'},
    {'Name': 'Your Name', 'ID': '123456', 'Department': 'Testing', 'Email': 'youremail@gmail.com'}
])
df.to_excel('examples/template_data.xlsx', index=False)
print("Created examples/template_data.xlsx")

# 2. Create Word Template
doc = Document()
doc.add_heading('Official Admit Card', 0)
doc.add_paragraph('Confidential Document')
doc.add_paragraph('Date: 2026-01-01')

p = doc.add_paragraph('Dear ')
p.add_run('<Name>').bold = True
p.add_run(',')

doc.add_paragraph('We are pleased to inform you that your profile has been approved.')

doc.add_paragraph('Details:', style='List Bullet')
p_id = doc.add_paragraph('Employee ID: ')
p_id.add_run('<ID>').bold = True

p_dept = doc.add_paragraph('Department: ')
p_dept.add_run('<Department>').bold = True

doc.add_paragraph('Please bring this document with you.')
doc.save('examples/template_document.docx')
print("Created examples/template_document.docx")
