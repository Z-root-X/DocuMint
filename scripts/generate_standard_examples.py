import pandas as pd
from docx import Document
import os

def create_examples():
    output_dir = "examples"
    os.makedirs(output_dir, exist_ok=True)
    
    # 1. Create Standard Excel Data
    df = pd.DataFrame([
        {
            "ID": "1001", 
            "Name": "Alice Johnson", 
            "Department": "Computer Science", 
            "Email": "alice@example.com"
        },
        {
            "ID": "1002", 
            "Name": "Bob Smith", 
            "Department": "Mathematics", 
            "Email": "bob@example.com"
        }
    ])
    excel_path = os.path.join(output_dir, "student_data.xlsx")
    df.to_excel(excel_path, index=False)
    print(f"Created {excel_path}")

    # 2. Create Standard Word Template
    doc = Document()
    doc.add_heading('University Admit Card', 0)
    
    p = doc.add_paragraph('Student Name: ')
    p.add_run('<Name>').bold = True
    
    p = doc.add_paragraph('Student ID: ')
    p.add_run('<ID>').bold = True
    
    p = doc.add_paragraph('Department: ')
    p.add_run('<Department>').bold = True
    
    doc.add_paragraph('This is an important document. Please present it at the exam hall.')
    
    doc_path = os.path.join(output_dir, "admit_card_template.docx")
    doc.save(doc_path)
    print(f"Created {doc_path}")

if __name__ == "__main__":
    create_examples()
