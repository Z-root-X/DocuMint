import os
import re
import time
import pandas as pd
from docx import Document
from datetime import datetime
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, Callable, List
import logging

import smtplib
from email.message import EmailMessage
import mimetypes
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Abstract Base Classes (The "Universal" Layer) ---

class DocumentConverter(ABC):
    """Abstract base class for converting documents to PDF."""
    
    @abstractmethod
    def convert_to_pdf(self, input_path: str, output_path: str) -> None:
        """Converts a source document to PDF."""
        pass

class EmailSender(ABC):
    """Abstract base class for sending emails."""
    
    @abstractmethod
    def send_email(self, to_email: str, subject: str, body_html: str, attachments: List[str] = []) -> None:
        """Sends an email with optional attachments."""
        pass

# --- Windows & Universal Implementations ---

class UniversalDocumentConverter(DocumentConverter):
    """Universal converter that dynamically selects the best available engine (Word COM, docx2pdf, or LibreOffice)."""
    
    def convert_to_pdf(self, input_path: str, output_path: str) -> None:
        abs_input = os.path.abspath(input_path)
        abs_output = os.path.abspath(output_path)
        
        # 1. Try Windows Word COM
        if os.name == 'nt':
            try:
                import win32com.client
                word_app = None
                doc = None
                try:
                    word_app = win32com.client.Dispatch("Word.Application")
                    word_app.Visible = False
                    doc = word_app.Documents.Open(abs_input)
                    doc.SaveAs(abs_output, FileFormat=17) # 17 = wdFormatPDF
                    return
                finally:
                    if doc:
                        doc.Close(False)
                    if word_app:
                        word_app.Quit()
            except Exception as e:
                logger.warning(f"WinWordConverter unavailable or failed: {e}. Attempting fallbacks...")

        # 2. Try docx2pdf
        try:
            from docx2pdf import convert
            convert(abs_input, abs_output)
            return
        except Exception:
            pass

        # 3. Try LibreOffice CLI
        import subprocess
        output_dir = os.path.dirname(abs_output)
        for cmd in ['libreoffice', 'soffice']:
            try:
                subprocess.run(
                    [cmd, '--headless', '--convert-to', 'pdf', abs_input, '--outdir', output_dir],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=30
                )
                generated_pdf = os.path.splitext(abs_input)[0] + '.pdf'
                if os.path.exists(generated_pdf) and generated_pdf != abs_output:
                    os.replace(generated_pdf, abs_output)
                return
            except Exception:
                continue

        raise RuntimeError(
            "No PDF conversion engine available. Please install Microsoft Word (Windows) or LibreOffice (Linux/Mac)."
        )

class WinWordConverter(UniversalDocumentConverter):
    """Compatibility alias for UniversalDocumentConverter."""
    pass

class WinOutlookSender(EmailSender):
    """Uses Microsoft Outlook (via COM) to send emails."""
    
    def __init__(self):
        import win32com.client
        self.outlook = win32com.client.Dispatch("Outlook.Application")

    def send_email(self, to_email: str, subject: str, body_html: str, attachments: List[str] = []) -> None:
        try:
            mail = self.outlook.CreateItem(0)
            mail.To = to_email
            mail.Subject = subject
            mail.HTMLBody = body_html
            
            for att in attachments:
                if os.path.exists(att):
                    mail.Attachments.Add(os.path.abspath(att))
            
            mail.Send()
        except Exception as e:
            logger.error(f"Outlook send failed: {e}")
            raise e

class SMTPSender(EmailSender):
    """Uses Python's smtplib to send emails."""
    
    def __init__(self, host: str, port: int, user: str, password: str):
        self.host = host
        self.port = port
        self.user = user
        self.password = password

    def send_email(self, to_email: str, subject: str, body_html: str, attachments: List[str] = []) -> None:
        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = self.user
        msg['To'] = to_email
        msg.set_content("Please enable HTML to view this email.")
        msg.add_alternative(body_html, subtype='html')

        for att in attachments:
            if os.path.exists(att):
                ctype, encoding = mimetypes.guess_type(att)
                if ctype is None or encoding is not None:
                    # No guess could be made, or the file is encoded (compressed), so
                    # use a generic bag-of-bits type.
                    ctype = 'application/octet-stream'
                
                maintype, subtype = ctype.split('/', 1)
                
                with open(att, 'rb') as f:
                    file_data = f.read()
                    msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=os.path.basename(att))

        try:
            with smtplib.SMTP_SSL(self.host, self.port) as smtp:
                smtp.login(self.user, self.password)
                smtp.send_message(msg)
        except Exception as e:
            logger.error(f"SMTP send failed: {e}")
            raise e

# --- Core Business Logic ---

def replace_placeholders_in_doc(doc: Document, replacements: Dict[str, Any]) -> None:
    """Replaces placeholders in a .docx document."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

def _replace_in_paragraph(paragraph, replacements: Dict[str, Any]) -> None:
    """Helper to replace placeholders in a paragraph while preserving run formatting when possible."""
    expanded_replacements = {}
    for k, v in replacements.items():
        val_str = str(v)
        expanded_replacements[k] = val_str
        clean_k = str(k).strip('<>{}[]')
        expanded_replacements[f"<{clean_k}>"] = val_str
        expanded_replacements[f"{{{{{clean_k}}}}}"] = val_str
        expanded_replacements[f"[{clean_k}]"] = val_str

    full_text = "".join(run.text for run in paragraph.runs)
    if not any(key in full_text for key in expanded_replacements):
        return

    sorted_keys = sorted(expanded_replacements.keys(), key=len, reverse=True)

    # 1. First attempt: run-by-run replacement to preserve bold/italic/font formatting
    for run in paragraph.runs:
        for key in sorted_keys:
            val = expanded_replacements[key]
            if key in run.text:
                run.text = run.text.replace(key, val)

    # 2. Fallback: if placeholder was split across runs by Word editor
    remaining_text = "".join(run.text for run in paragraph.runs)
    if any(key in remaining_text for key in sorted_keys):
        new_text = remaining_text
        for key in sorted_keys:
            val = expanded_replacements[key]
            new_text = new_text.replace(key, val)
        if new_text != remaining_text:
            p = paragraph._p
            for child in list(p):
                p.remove(child)
            paragraph.add_run(new_text)

def is_valid_email(email_address: str) -> bool:
    """Validates an email address."""
    pattern = r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$"
    return re.match(pattern, email_address) is not None

def validate_placeholders(data_file: str, template_file: str) -> tuple[bool, list[str]]:
    """
    Checks if all placeholders in the template exist as columns in the data file.
    Supports <Column>, {{Column}}, and [Column] syntax.
    Returns (True, []) if valid, or (False, list_of_missing_columns).
    """
    if not os.path.exists(data_file) or not os.path.exists(template_file):
        return False, ["File not found"]

    try:
        df = pd.read_excel(data_file) if data_file.endswith(('.xlsx', '.xls')) else pd.read_csv(data_file)
        df.columns = df.columns.astype(str).str.strip()
        data_columns = set(df.columns)

        doc = Document(template_file)
        placeholders = set()
        pattern = r"(?:<|{{|\[)([^>}\]]+)(?:>|}}|\])"
        
        def extract_from_text(text):
            matches = re.findall(pattern, text)
            for m in matches:
                placeholders.add(m.strip())

        for para in doc.paragraphs:
            extract_from_text(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        extract_from_text(para.text)

        missing = [p for p in placeholders if p not in data_columns]
        return (len(missing) == 0), missing

    except Exception as e:
        logger.error(f"Validation failed: {e}")
        return False, [str(e)]


def process_emails(
    data_file: str, 
    template_file: str, 
    pdf_folder: str, 
    logs_folder: str, 
    log_callback: Callable[[str], None], 
    email_subject: str, 
    email_body: str, 
    pdf_filename_format: str, 
    retries: int, 
    delay: int, 
    dry_run: bool = False, 
    test_email: Optional[str] = None,
    email_config: Optional[Dict[str, Any]] = None
) -> None:
    
    log_file_path = os.path.join(logs_folder, "documint_log.xlsx")
    log_records = []

    # 1. Initialize Engines
    doc_converter = WinWordConverter() # Default for now, can be configured later
    email_sender = None

    if not dry_run:
        try:
            # Determine Email Sender based on config
            provider = email_config.get("provider", "outlook") if email_config else "outlook"
            
            if provider == "smtp":
                if not email_config:
                    raise ValueError("SMTP configuration missing")
                email_sender = SMTPSender(
                    host=email_config.get("smtp_host", ""),
                    port=int(email_config.get("smtp_port", 465)),
                    user=email_config.get("smtp_user", ""),
                    password=email_config.get("smtp_password", "")
                )
            else:
                email_sender = WinOutlookSender()
                
        except Exception as e:
            log_callback(f"❌ Email System Init Failed: {e}")
            return

    # 2. Load Data
    try:
        df = pd.read_excel(data_file)
        df.columns = df.columns.str.strip()
    except Exception as e:
        log_callback(f"❌ Error reading Excel file: {str(e)}")
        return

    # 3. Handle Test Mode
    if test_email:
        dummy_data = {col: f"Test {col}" for col in df.columns}
        dummy_data["Email"] = test_email
        df = pd.DataFrame([dummy_data])

    # 4. Processing Loop
    # 4. Processing Loop
    import concurrent.futures
    
    # Helper to process a single row
    def process_row(args):
        index, row = args
        local_logs = []
        try:
            email = str(row.get("Email", "")).strip()
            if not is_valid_email(email):
                log_callback(f"⚠️ Skipped invalid email: {email}")
                return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": "Failed: Invalid Email", "Timestamp": datetime.now()}]

            replacements = {f"<{col}>": str(val).strip() for col, val in row.items()}
            
            # A. Generate Word Doc
            filename = pdf_filename_format.format(**replacements)
            filename = "".join(c for c in filename if c.isalnum() or c in (' ', '_', '-'))
            
            word_path = os.path.join(pdf_folder, f"{filename}.docx")
            pdf_path = os.path.join(pdf_folder, f"{filename}.pdf")

            # Word operations must be main thread or carefully managed if using COM
            # Since we are using WinWordConverter (COM), we MUST be careful.
            # COM and Threads don't mix well without CoInitialize.
            # Strategy: Generate DOCX is pure python-docx (Safe).
            # Convert to PDF via COM (Unsafe in threads without care).
            
            doc = Document(template_file)
            replace_placeholders_in_doc(doc, replacements)
            doc.save(word_path)

            # B. Convert to PDF
            try:
                # If threading, we need CoInitialize for COM
                if isinstance(doc_converter, WinWordConverter):
                    import pythoncom
                    pythoncom.CoInitialize()
                    
                doc_converter.convert_to_pdf(word_path, pdf_path)
            except Exception as e:
                log_callback(f"❌ PDF Conv Error for {email}: {e}")
                return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": f"Conversion Error: {e}", "Timestamp": datetime.now()}]
            finally:
                if isinstance(doc_converter, WinWordConverter):
                    import pythoncom
                    pythoncom.CoUninitialize()
                if os.path.exists(word_path):
                    os.remove(word_path)

            # C. Send Email
            if not dry_run and email_sender:
                sent = False
                
                final_body = email_body
                final_subject = email_subject
                for k, v in replacements.items():
                    final_body = final_body.replace(k, str(v))
                    final_subject = final_subject.replace(k, str(v))

                for attempt in range(retries + 1):
                    try:
                        email_sender.send_email(email, final_subject, final_body, [pdf_path])
                        sent = True
                        break
                    except Exception as e:
                        if attempt < retries:
                            time.sleep(delay)
                        else:
                            raise e
                
                if sent:
                    log_callback(f"✅ Sent to {email}")
                    return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": "Success", "Timestamp": datetime.now()}]
            else:
                log_callback(f"📝 Generated PDF for {email} (Dry Run)")
                return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": "Dry Run", "Timestamp": datetime.now()}]

            time.sleep(delay)
            return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": "Success", "Timestamp": datetime.now()}]

        except Exception as e:
            log_callback(f"❌ Error processing row {index}: {e}")
            return [{"Name": str(row.get("Name", "N/A")).strip(), "Email": email, "Status": f"Error: {e}", "Timestamp": datetime.now()}]

    # Execution Strategy
    # If SMTP, we can use threads (faster).
    # If Outlook, we MUST use main thread (serial) or single threaded due to COM.
    
    use_threading = False
    if email_config and email_config.get("provider") == "smtp":
        use_threading = True
        
    rows = [(i, r) for i, r in df.iterrows()]
    
    if use_threading:
        log_callback("🚀 Speed Mode: Parallel Execution Enabled (SMTP)")
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            results = executor.map(process_row, rows)
            for res in results:
                log_records.extend(res)
    else:
        log_callback("⚠️ Standard Mode: Serial Execution (Outlook requires this)")
        for args in rows:
            res = process_row(args)
            log_records.extend(res)

    # 5. Save Logs

    # 5. Save Logs
    try:
        new_log_df = pd.DataFrame(log_records)
        if os.path.exists(log_file_path):
            old_log_df = pd.read_excel(log_file_path)
            new_log_df = pd.concat([old_log_df, new_log_df], ignore_index=True)
        new_log_df.to_excel(log_file_path, index=False)
        log_callback(f"📄 Log saved to {log_file_path}")
    except Exception as e:
        log_callback(f"❌ Failed to save log: {e}")

def _log_status(records: list, row: pd.Series, email: str, status: str):
    records.append({
        "Name": str(row.get("Name", "N/A")).strip(),
        "Email": email,
        "Status": status,
        "Timestamp": datetime.now()
    })